from pathlib import Path

from docx import Document
from openpyxl import Workbook

from fastapi.responses import FileResponse


EXPORT_DIR = Path("exports")

EXPORT_DIR.mkdir(
    exist_ok=True
)


def export_docx(
    data
):

    question = data.get(
        "question",
        ""
    )

    answer = data.get(
        "answer",
        ""
    )

    path = (
        EXPORT_DIR
        / "answer.docx"
    )

    doc = Document()

    doc.add_heading(
        "AI Knowledge Assistant",
        level=1
    )

    doc.add_heading(
        "Question",
        level=2
    )

    doc.add_paragraph(
        question
    )

    doc.add_heading(
        "Answer",
        level=2
    )

    doc.add_paragraph(
        answer
    )

    doc.save(
        str(path)
    )

    return FileResponse(
        str(path),
        filename="answer.docx"
    )


def export_xlsx(
    data
):

    question = data.get(
        "question",
        ""
    )

    answer = data.get(
        "answer",
        ""
    )

    path = (
        EXPORT_DIR
        / "answer.xlsx"
    )

    wb = Workbook()

    ws = wb.active

    ws["A1"] = "Question"
    ws["B1"] = "Answer"

    ws["A2"] = question
    ws["B2"] = answer

    wb.save(
        str(path)
    )

    return FileResponse(
        str(path),
        filename="answer.xlsx"
    )